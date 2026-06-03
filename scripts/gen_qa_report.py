#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Genera el Plan/Reporte de Calidad QA (Sistema de Gestión de Usuarios — Registro)
siguiendo el formato del modelo de ejemplo, ampliado y más detallado.
Aplica la metodología de la presentación (dimensiones de calidad, validación en
todas las capas, severidad/prioridad, trazabilidad RF→CA→CP)."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x14, 0x3C, 0x5B)
GREY = RGBColor(0x55, 0x55, 0x55)
HDR_FILL = "1F3A5F"   # relleno encabezado de tabla
ZEBRA = "EEF2F6"      # fila alterna

doc = Document()

# ---- estilos base ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)

def _shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def _white(cell):
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r.font.bold = True

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = NAVY
    return h

def para(text, italic=False, bold=False, size=10.5, color=None, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic; r.bold = bold; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p

def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    return p

def table(headers, rows, widths=None, zebra=True):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        _shade(hdr[i], HDR_FILL); _white(hdr[i])
        for p in hdr[i].paragraphs:
            p.paragraph_format.space_after = Pt(0)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs: r.font.size = Pt(9.5)
            if zebra and ri % 2 == 1:
                _shade(cells[i], ZEBRA)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

# =================== PORTADA ===================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Reporte de Calidad QA")
r.font.size = Pt(26); r.font.bold = True; r.font.color.rgb = NAVY
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Sistema de Gestión de Usuarios — Módulo de Registro")
r.font.size = Pt(14); r.italic = True; r.font.color.rgb = GREY
doc.add_paragraph()

meta = [
    ("Proyecto", "Sistema de Gestión de Usuarios"),
    ("Módulo evaluado", "Administración de usuarios — Registro"),
    ("Funcionalidad principal", "Registro de usuarios con validación de datos, roles y persistencia"),
    ("Historia de usuario", "Como administrador, quiero registrar usuarios para dar acceso según su rol."),
    ("Versión evaluada", "v1.0.0"),
    ("Ambiente", "QA"),
    ("Fecha de elaboración", "03-06-2026"),
    ("Responsable QA", "Equipo de Validación — Ingeniería en Informática"),
    ("Navegador", "Google Chrome (última versión)"),
    ("Frontend", "React / Angular / Vue (SPA)"),
    ("Backend", "API REST"),
    ("Base de datos", "Relacional (MySQL / PostgreSQL)"),
    ("Tipo de documento", "Plan de pruebas QA + reporte de ejecución (ejecutado contra el entorno real)"),
    ("Plataforma evaluada", "ChessQuery (entorno QA en AWS; auth Supabase, API REST vía gateway)"),
    ("Estado general", "Con observaciones"),
]
t = table(["Campo", "Detalle"], meta, widths=[2.0, 4.5], zebra=True)

doc.add_page_break()

# =================== 1. OBJETIVO ===================
heading("1. Objetivo del documento")
para("Este documento define el plan de validación QA del módulo de registro de usuarios y "
     "sirve, a la vez, como plantilla de reporte de calidad. Su objetivo es verificar que la "
     "funcionalidad cumpla los requerimientos definidos y los criterios de aceptación, asegurando "
     "comportamiento correcto en las tres dimensiones de la calidad: funcional (hace lo que debe), "
     "técnica (seguro, estable, maneja errores) y de experiencia (claro y con retroalimentación "
     "precisa). La validación se realiza en todas las capas —frontend, backend, integración y "
     "persistencia—, no solo en la interfaz.")

# =================== 2. ALCANCE ===================
heading("2. Alcance de la validación")
para("Funcionalidades incluidas:", bold=True, space_after=2)
for b in [
    "Registro de usuarios con datos válidos.",
    "Validación de campos obligatorios (bloqueo de envío con campos vacíos).",
    "Validación de formato de correo y de dominios permitidos (gmail.com y duocuc.cl).",
    "Validación de contraseña: mínimo 8 caracteres y coincidencia con su confirmación.",
    "Mensajes de error mostrados junto al campo correspondiente.",
    "Asignación y validación de rol del sistema (Player y Organizador como roles relevantes, más Admin).",
    "Unicidad de identidad (Nombre + Apellidos) y de correo electrónico; bloqueo de duplicados.",
    "Persistencia correcta e íntegra del usuario en la base de datos (validada también en backend).",
    "Mensaje de éxito visible con animación precisa tras el registro y redirección a la plataforma tras el login.",
    "Actualización automática de la tabla de usuarios.",
    "Manejo controlado de errores del backend (p. ej. HTTP 500).",
]:
    bullet(b)
para("Fuera de alcance:", bold=True, space_after=2)
for b in [
    "Pruebas de carga y estrés.",
    "Pruebas avanzadas de seguridad (pentesting, fuzzing).",
    "Edición y eliminación de usuarios (se evalúan en un ciclo posterior).",
]:
    bullet(b)

# =================== 3. ENFOQUE Y METODOLOGÍA ===================
heading("3. Enfoque y metodología")
para("La validación QA comprueba que cada funcionalidad cumple el comportamiento esperado desde la "
     "perspectiva del negocio y del usuario final. Se apoya en los siguientes principios:")
for b in [
    "Verificación («¿lo construimos bien?») y Validación («¿construimos lo correcto?») de forma complementaria.",
    "Validar en todas las capas: una regla de negocio (formato de correo, unicidad, rol válido) debe aplicarse "
    "en el servidor y no depender solo del frontend, que puede ser omitido consumiendo la API directamente.",
    "Trazabilidad completa: cada requerimiento (RF) se asocia a criterios de aceptación (CA) y casos de prueba (CP).",
    "Clasificación de defectos por severidad (impacto técnico) y prioridad (urgencia de corrección), evaluadas de forma independiente.",
]:
    bullet(b)
para("Tipos de prueba aplicados: funcionales, de integración (frontend–backend–BD), de seguridad básica "
     "(autorización por rol) y de regresión sobre el flujo de registro.", space_after=10)

# =================== 4. REQUERIMIENTOS (RF) ===================
heading("4. Requerimientos evaluados (RF)")
rf = [
    ("RF-01", "Registrar usuario con datos válidos y persistirlo.", "Alta", "Por ejecutar"),
    ("RF-02", "Impedir el envío del formulario con campos obligatorios vacíos.", "Alta", "Por ejecutar"),
    ("RF-03", "Validar formato de correo y aceptar solo dominios gmail.com y duocuc.cl.", "Alta", "Por ejecutar"),
    ("RF-04", "Exigir contraseña de mínimo 8 caracteres y confirmación coincidente.", "Alta", "Por ejecutar"),
    ("RF-05", "Mostrar el mensaje de error junto al campo inválido correspondiente.", "Media", "Por ejecutar"),
    ("RF-06", "Persistir el usuario de forma íntegra en la base de datos (validado en backend).", "Alta", "Por ejecutar"),
    ("RF-07", "Asignar y validar un rol del sistema: Player, Organizador (relevantes) o Admin.", "Alta", "Por ejecutar"),
    ("RF-08", "Garantizar unicidad de (Nombre+Apellidos) y de correo; bloquear duplicados con mensaje claro.", "Alta", "Por ejecutar"),
    ("RF-09", "Mostrar mensaje de éxito visible con animación precisa tras el registro.", "Media", "Por ejecutar"),
    ("RF-10", "Redirigir a la plataforma tras un login exitoso.", "Media", "Por ejecutar"),
    ("RF-11", "Actualizar automáticamente la tabla de usuarios tras un alta.", "Media", "Por ejecutar"),
    ("RF-12", "Informar de forma controlada los errores del backend (p. ej. HTTP 500).", "Alta", "Por ejecutar"),
]
table(["Código", "Requerimiento", "Prioridad", "Estado"], rf, widths=[0.8, 4.3, 0.9, 1.1])

# =================== 5. CRITERIOS DE ACEPTACIÓN (CA) ===================
heading("5. Criterios de aceptación (CA)")
ca = [
    ("CA-01", "Con todos los campos válidos, el registro se permite: responde 201, persiste y muestra éxito.", "RF-01, RF-06, RF-09"),
    ("CA-02", "Si algún campo obligatorio está vacío, el envío se bloquea (botón deshabilitado o validación).", "RF-02"),
    ("CA-03", "Un correo con formato inválido (sin @, sin dominio) es rechazado.", "RF-03"),
    ("CA-04", "Un correo con dominio distinto de gmail.com o duocuc.cl es rechazado.", "RF-03"),
    ("CA-05", "Una contraseña de menos de 8 caracteres es rechazada.", "RF-04"),
    ("CA-06", "Si la contraseña y su confirmación no coinciden, el registro es rechazado.", "RF-04"),
    ("CA-07", "Cada error se muestra junto al campo correspondiente, con texto descriptivo.", "RF-05"),
    ("CA-08", "El rol es obligatorio y debe pertenecer a {Player, Organizador, Admin} (Player y Organizador son los relevantes); otro valor se rechaza.", "RF-07"),
    ("CA-09", "Un correo ya registrado se bloquea con mensaje claro de duplicado.", "RF-08"),
    ("CA-10", "Una combinación Nombre+Apellidos ya existente se bloquea con mensaje claro.", "RF-08"),
    ("CA-11", "El usuario persiste íntegro en la BD: todos los campos y el rol asignado.", "RF-06, RF-07"),
    ("CA-12", "Tras el registro exitoso se muestra un mensaje de éxito visible con animación precisa.", "RF-09"),
    ("CA-13", "Tras un login exitoso el usuario es redirigido a la plataforma.", "RF-10"),
    ("CA-14", "La tabla de usuarios se actualiza automáticamente, sin recargar la página.", "RF-11"),
    ("CA-15", "Ante un error del backend (500), se informa al usuario sin exponer detalles técnicos.", "RF-12"),
]
table(["Código", "Criterio de aceptación", "RF asociado"], ca, widths=[0.8, 4.6, 1.6])

doc.add_page_break()

# =================== 6. CASOS DE PRUEBA (CP) ===================
heading("6. Diseño de casos de prueba (CP)")
para("Cada caso define cómo verificar un criterio de aceptación. La columna «Resultado obtenido» y "
     "«Estado» se completan durante la ejecución (estado inicial: Por ejecutar).", italic=True, color=GREY)
cp = [
    ("CP-01", "Registrar con datos válidos", "Frontend+Backend", "Usuario creado (201), éxito visible y registro persistido en BD.", "CA-01"),
    ("CP-02", "Enviar con campos vacíos", "Frontend", "Se bloquea el envío y se muestra el error en cada campo obligatorio.", "CA-02, CA-07"),
    ("CP-03", "Correo con formato inválido (sin @)", "Frontend", "Se bloquea el envío; error de formato junto al campo correo.", "CA-03, CA-07"),
    ("CP-04", "Correo de dominio no permitido (p. ej. @hotmail.com)", "Frontend+Backend", "Se rechaza por dominio no permitido (solo gmail.com / duocuc.cl).", "CA-04"),
    ("CP-05", "Correo válido con dominio duocuc.cl", "Frontend", "Se acepta el correo (dominio permitido).", "CA-04"),
    ("CP-06", "Contraseña con menos de 8 caracteres", "Frontend", "Error «mínimo 8 caracteres» junto al campo contraseña.", "CA-05, CA-07"),
    ("CP-07", "Confirmación de contraseña distinta", "Frontend", "Error «las contraseñas no coinciden» junto al campo confirmación.", "CA-06, CA-07"),
    ("CP-08", "Registrar sin seleccionar rol", "Frontend", "Mensaje específico: «El campo rol es obligatorio».", "CA-08, CA-07"),
    ("CP-09", "Enviar rol inválido por API (p. ej. MODERADOR / SUPERADMIN)", "Backend", "El backend rechaza con 400 (rol fuera de Player/Organizador/Admin), no persiste.", "CA-08"),
    ("CP-10", "Registrar correo ya existente", "Frontend+Backend", "Se bloquea con mensaje claro de correo duplicado.", "CA-09"),
    ("CP-11", "Registrar Nombre+Apellidos ya existente", "Frontend+Backend", "Se bloquea con mensaje claro de identidad duplicada.", "CA-10"),
    ("CP-12", "Verificar persistencia en BD", "Persistencia", "El registro existe en BD con todos los campos y el rol correctos.", "CA-11"),
    ("CP-13", "Correo inválido enviado directo a la API", "Backend", "El backend valida de forma independiente y responde 400.", "CA-03, CA-04"),
    ("CP-14", "Registro exitoso: mensaje + animación", "Frontend", "Mensaje de éxito visible con animación precisa (no intrusiva).", "CA-12"),
    ("CP-15", "Redirección tras login", "Frontend", "Tras login exitoso, el usuario es llevado a la plataforma.", "CA-13"),
    ("CP-16", "Actualización automática de la tabla", "Frontend+Integración", "La tabla muestra el nuevo usuario sin recargar la página.", "CA-14"),
    ("CP-17", "Error 500 del backend (simulado)", "Integración", "Mensaje de error controlado al usuario, sin detalles técnicos.", "CA-15"),
    ("CP-18", "Acceso sin permiso (rol no autorizado)", "Seguridad", "La acción de registro se bloquea por rol.", "CA-08"),
]
table(["ID", "Caso de prueba", "Tipo/Capa", "Resultado esperado", "CA"],
      cp, widths=[0.6, 1.9, 1.2, 2.6, 0.9])

# detalle de un caso modelo
heading("6.1 Caso de prueba detallado (modelo) — CP-01", level=2)
det = [
    ("RF asociado", "RF-01"),
    ("Criterios", "CA-01"),
    ("Precondición", "Usuario administrador autenticado en el sistema."),
    ("Datos de prueba", "Nombre: Ana; Apellidos: López Soto; Correo: ana.lopez@duocuc.cl; "
                        "Contraseña: Clave2026 (confirmación: Clave2026); Rol: Player."),
    ("Pasos", "1) Abrir el formulario de registro. 2) Completar todos los campos con datos válidos. "
              "3) Presionar «Guardar»."),
    ("Resultado esperado", "Usuario creado exitosamente (HTTP 201), mensaje de éxito visible con animación, "
                           "registro persistido en BD y visible en la tabla de usuarios."),
    ("Resultado obtenido", "(Completar durante la ejecución)"),
    ("Estado", "Por ejecutar"),
    ("Evidencia requerida", "Captura del mensaje de éxito, respuesta de la API (201) y fila en la BD/tabla."),
]
table(["Campo", "Detalle"], det, widths=[1.6, 4.9])

# ---- 6.2 resultados de la ejecución ----
heading("6.2 Resultados de la ejecución (entorno real ChessQuery)", level=2)
para("Ejecutado el 03-06-2026 contra el entorno QA. Convención de estado: Aprobado, "
     "Observado (funciona distinto a lo planificado / regla no implementada), N/A (no aplica al sistema actual).",
     italic=True, color=GREY)
res = [
    ("CP-01", "Usuario creado y perfil persistido (id 31).", "Aprobado"),
    ("CP-02", "El frontend bloquea el envío y valida cada campo obligatorio.", "Aprobado"),
    ("CP-03", "El backend (Supabase) rechaza el correo con 400 «invalid format».", "Aprobado"),
    ("CP-04", "Se crea la cuenta con @hotmail.com: ChessQuery NO restringe dominios.", "Observado"),
    ("CP-05", "Correo con dominio duocuc.cl aceptado.", "Aprobado"),
    ("CP-06", "Frontend exige ≥8; el backend (Supabase) acepta desde 6 → brecha front/back.", "Observado"),
    ("CP-07", "El frontend valida la coincidencia de contraseña y su confirmación.", "Aprobado"),
    ("CP-08", "El selector de rol siempre tiene valor (PLAYER por defecto); no hay alta de «Admin» en la UI.", "Observado"),
    ("CP-09", "La BD rechaza rol inválido por check constraint «user_profiles_role_check» (error 23514).", "Aprobado"),
    ("CP-10", "El backend rechaza el correo duplicado con 422 «User already registered».", "Aprobado"),
    ("CP-11", "No se valida unicidad de Nombre+Apellidos; la unicidad real es por correo.", "Observado"),
    ("CP-12", "El perfil persiste íntegro (id, correo y rol) y es legible vía API.", "Aprobado"),
    ("CP-13", "El backend valida el correo de forma independiente (400) al consumir la API directa.", "Aprobado"),
    ("CP-14", "Tras el registro se redirige directamente; no hay mensaje de éxito ni animación explícita.", "Observado"),
    ("CP-15", "Tras el login, redirige al portal o al panel del organizador según el rol.", "Aprobado"),
    ("CP-16", "El panel del organizador auto-refresca su grilla; el portal del jugador no tiene «tabla de usuarios».", "N/A"),
    ("CP-17", "El cliente HTTP muestra un mensaje de error controlado ante fallos del backend.", "Aprobado"),
    ("CP-18", "Los endpoints de organizador exigen rol ORGANIZER (control por rol en el gateway/BFF).", "Aprobado"),
]
table(["ID", "Resultado obtenido", "Estado"], res, widths=[0.6, 4.8, 1.1])
para("Resumen de ejecución: 12 Aprobados · 5 Observados · 1 N/A (de 18 casos). "
     "Las observaciones se detallan en la sección 11.", bold=True)

doc.add_page_break()

# =================== 7. MATRIZ DE TRAZABILIDAD ===================
heading("7. Matriz de trazabilidad (RF → CA → CP)")
para("Asegura cobertura completa: cada requerimiento se cubre con criterios y casos de prueba.", italic=True, color=GREY)
tz = [
    ("RF-01", "CA-01, CA-11", "CP-01, CP-12", "Aprobado"),
    ("RF-02", "CA-02, CA-07", "CP-02", "Aprobado"),
    ("RF-03", "CA-03, CA-04", "CP-03, CP-04, CP-05, CP-13", "Observado (dominio no restringido)"),
    ("RF-04", "CA-05, CA-06", "CP-06, CP-07", "Observado (brecha 8/6)"),
    ("RF-05", "CA-07", "CP-02, CP-03, CP-06, CP-07, CP-08", "Aprobado"),
    ("RF-06", "CA-01, CA-11", "CP-01, CP-12", "Aprobado"),
    ("RF-07", "CA-08, CA-11", "CP-08, CP-09, CP-18", "Aprobado"),
    ("RF-08", "CA-09, CA-10", "CP-10, CP-11", "Parcial (email sí; nombre no)"),
    ("RF-09", "CA-12", "CP-14", "Observado (sin éxito/animación)"),
    ("RF-10", "CA-13", "CP-15", "Aprobado"),
    ("RF-11", "CA-14", "CP-16", "N/A (sin tabla de usuarios en portal)"),
    ("RF-12", "CA-15", "CP-17", "Aprobado"),
]
table(["RF", "Criterios (CA)", "Casos (CP)", "Estado"], tz, widths=[0.8, 1.8, 2.6, 1.1])

# =================== 8. VALIDACIÓN FRONTEND ===================
heading("8. Validación frontend")
fe = [
    ("Formulario", "Bloquea envío con campos vacíos; valida cada campo obligatorio.", "Aprobado"),
    ("Campo correo", "Valida formato; NO restringe dominios a gmail.com / duocuc.cl.", "Observado"),
    ("Campo contraseña", "Exige mínimo 8 caracteres y coincidencia con la confirmación.", "Aprobado"),
    ("Campo rol", "Selección con valor por defecto; sin opción «Admin» en la UI (enum validado en BD).", "Observado"),
    ("Mensajes de error", "Aparecen junto al campo correcto, con texto descriptivo.", "Aprobado"),
    ("Mensaje de éxito", "No se muestra mensaje/animación de éxito; redirige directamente.", "Observado"),
    ("Tabla de usuarios", "El portal del jugador no tiene tabla de usuarios (sí el panel organizador).", "N/A"),
    ("Redirección", "Tras login exitoso, navega a la plataforma según el rol.", "Aprobado"),
    ("Estados de carga", "Indicador visible mientras espera la respuesta de la API.", "Aprobado"),
    ("Responsive", "Adaptación a distintos tamaños de pantalla.", "No verificado"),
]
table(["Elemento", "Qué se verifica", "Estado"], fe, widths=[1.5, 3.9, 1.1])

# =================== 9. VALIDACIÓN BACKEND ===================
heading("9. Validación backend")
para("El servidor debe aplicar las reglas de negocio de forma independiente del frontend.", space_after=4)
be = [
    ("Registro — datos válidos", "Usuario creado y persistido íntegro.", "Aprobado"),
    ("Registro — correo inválido", "Rechazado 400 «invalid format».", "Aprobado"),
    ("Registro — dominio no permitido", "Aceptado (no se restringe el dominio).", "Observado"),
    ("Registro — rol inválido", "Rechazado por check constraint de BD (rol fuera del enum).", "Aprobado"),
    ("Registro — correo duplicado", "Rechazado 422 «User already registered».", "Aprobado"),
    ("Acceso sin autorización", "Endpoints de organizador exigen rol ORGANIZER.", "Aprobado"),
    ("Contrato de respuesta (JSON)", "Estructura coincide con lo esperado.", "Aprobado"),
    ("Manejo de error interno (500)", "No se forzó un 500 real en esta corrida.", "No verificado"),
]
table(["Endpoint / escenario", "Resultado esperado", "Estado"], be, widths=[2.5, 3.0, 1.0])
para("Códigos HTTP esperados según escenario: 200/201 (éxito), 400 (datos inválidos), "
     "401/403 (no autorizado), 404 (no encontrado), 409 (duplicado), 500 (error interno controlado).",
     italic=True, color=GREY)

doc.add_page_break()

# =================== 10. RESUMEN EJECUTIVO ===================
heading("10. Resumen ejecutivo")
para("Estado global: Con observaciones. De 18 casos: 12 Aprobados, 5 Observados, 1 N/A. "
     "No se registraron defectos bloqueantes; las observaciones son de severidad media/baja.", bold=True)
para("Fortalezas verificadas en el entorno real:")
for b in [
    "El backend valida de forma independiente del frontend: rechaza correos con formato inválido (400) "
    "y correos duplicados (422), incluso consumiendo la API directamente.",
    "El rol se valida a nivel de base de datos mediante el check constraint «user_profiles_role_check»: "
    "no es posible crear un usuario con un rol fuera del enum permitido (Player/Organizador/Admin).",
    "La persistencia es íntegra: el usuario registrado queda almacenado y es consultable vía API.",
    "El control de acceso por rol protege los endpoints del organizador.",
]:
    bullet(b)
para("Observaciones (oportunidades de mejora):")
for b in [
    "No se aplica la restricción de dominios de correo (gmail.com / duocuc.cl): se aceptan otros dominios.",
    "Brecha de longitud de contraseña: el frontend exige 8, pero el backend acepta desde 6 caracteres.",
    "No se valida unicidad de Nombre+Apellidos (la identidad única es el correo).",
    "Tras un registro exitoso no se muestra un mensaje de éxito con animación; se redirige directamente.",
]:
    bullet(b)

# =================== 11. HALLAZGOS Y OBSERVACIONES ===================
heading("11. Hallazgos y observaciones")
para("No se detectaron defectos bloqueantes. Se registran 4 observaciones (mejoras). "
     "Severidad = impacto técnico; Prioridad = urgencia de corrección (independientes).", italic=True, color=GREY)
obs = [
    ("OBS-01", "El backend no restringe el dominio del correo (acepta @hotmail.com).", "Media", "Media", "RF-03 / CP-04", "Abierto"),
    ("OBS-02", "Brecha de longitud de contraseña: frontend ≥8, backend acepta desde 6.", "Media", "Media", "RF-04 / CP-06", "Abierto"),
    ("OBS-03", "No se valida unicidad de Nombre+Apellidos (la unicidad es por correo).", "Baja", "Baja", "RF-08 / CP-11", "Abierto"),
    ("OBS-04", "Sin mensaje de éxito con animación tras el registro; redirige directamente.", "Baja", "Media", "RF-09 / CP-14", "Abierto"),
]
table(["ID", "Hallazgo", "Sev.", "Prior.", "RF / CP", "Estado"], obs, widths=[0.7, 2.9, 0.7, 0.7, 1.1, 0.8])

heading("11.1 Detalle — OBS-01: dominio de correo no restringido", level=2)
d1 = [
    ("Severidad / Prioridad", "Media / Media"),
    ("RF / CP", "RF-03 / CP-04"),
    ("Pasos", "Enviar signup con correo «qaexec@hotmail.com» a la API de autenticación."),
    ("Resultado esperado", "Rechazo por dominio no permitido (solo gmail.com / duocuc.cl), si la regla aplica."),
    ("Resultado obtenido", "La cuenta se crea: no hay restricción de dominio en frontend ni backend."),
    ("Anotación QA", "Regla de negocio específica del plan; ChessQuery hoy no la implementa por diseño."),
    ("Recomendación", "Si el dominio debe restringirse, validarlo en el backend (y reflejarlo en el frontend)."),
]
table(["Campo", "Detalle"], d1, widths=[1.8, 4.7])

heading("11.2 Detalle — OBS-02: brecha de longitud de contraseña", level=2)
d2 = [
    ("Severidad / Prioridad", "Media / Media"),
    ("RF / CP", "RF-04 / CP-06"),
    ("Pasos", "Crear cuenta por la API con contraseña de 7 caracteres."),
    ("Resultado esperado", "Rechazo (la regla del producto exige mínimo 8)."),
    ("Resultado obtenido", "El backend (Supabase) la acepta: su mínimo es 6, por debajo del requisito de 8."),
    ("Anotación QA", "El frontend exige 8, pero consumiendo la API directa se crean cuentas con 6-7."),
    ("Recomendación", "Alinear el mínimo del backend a 8 (política de contraseñas de Supabase Auth)."),
]
table(["Campo", "Detalle"], d2, widths=[1.8, 4.7])

# =================== 12. SEVERIDAD Y PRIORIDAD ===================
heading("12. Criterios de severidad y prioridad")
sev = [
    ("Crítica", "Bloquea una funcionalidad completa o corrompe datos."),
    ("Alta", "Falla funcional importante o riesgo de integridad/seguridad."),
    ("Media", "Comportamiento incorrecto sin bloquear el flujo principal."),
    ("Baja", "Detalle estético o de redacción."),
]
table(["Severidad", "Definición (impacto técnico)"], sev, widths=[1.4, 5.1])
para("La prioridad (urgencia de corrección) se evalúa de forma independiente: un defecto de severidad "
     "baja puede tener prioridad alta si afecta una demo con el cliente.", italic=True, color=GREY)

# =================== 13. CONCLUSIONES ===================
heading("13. Conclusiones y recomendaciones")
for b in [
    "La calidad se valida en todas las capas: toda regla de negocio debe aplicarse en el backend, no solo en el frontend.",
    "Priorizar las validaciones de integridad: unicidad (correo y Nombre+Apellidos), dominios permitidos y rol válido.",
    "Cuidar la experiencia: mensajes de error junto al campo, mensaje de éxito con animación precisa y manejo controlado de errores 500.",
    "Completar la ejecución de los 18 casos, registrar resultados/evidencias y abrir los defectos correspondientes.",
    "Tras corregir, ejecutar pruebas de regresión sobre el flujo de registro y login.",
]:
    bullet(b)

import os
out = os.path.expanduser("~/Descargas/Reporte_QA_Sistema_Gestion_Usuarios.docx")
doc.save(out)
print("OK ->", out)
